from docx import Document  # Importamos la librería para Word 

def extraer_texto_word(ruta_archivo):
    try:
        print(f"Abriendo documento Word: {ruta_archivo}...")
        
        # Cargamos el documento
        doc = Document(ruta_archivo)
        print(f"Word abierto - Contiene {len(doc.paragraphs)} párrafos detectados.")
        
        # Crea lista vacía para guardar los fragmentos de texto
        textos = []
        
        # Recorremos cada párrafo del documento
        for parrafo in doc.paragraphs:
            # Limpia los espacios en blanco extra
            texto_limpio = parrafo.text.strip()
            
            # Solo guardamos el párrafo si realmente tiene texto
            if texto_limpio:
                textos.append(texto_limpio)
                
        # Unimos todos los párrafos en un solo gran texto separado por saltos de línea (\n)
        texto_final = "\n".join(textos)
        
        print("\n--- INICIO DEL TEXTO EXTRAÍDO ---")
        print(texto_final)
        print("--- FIN DEL TEXTO EXTRAÍDO ---\n")
        
        return texto_final
        
    except Exception as e:
        print(f"Error al intentar leer el Word: {e}")
        return None

# Ejecutamos la función apuntando a nuestro archivo
ruta_de_prueba = "temp_files/prueba.docx"
extraer_texto_word(ruta_de_prueba)